from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Neues Dokument erstellen
doc = Document()

# Dokumentenstile anpassen
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Titel
title = doc.add_heading('GEHEIMHALTUNGSVEREINBARUNG (NDA)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Zwischen
doc.add_heading('zwischen', level=1)

# Erste Partei
p = doc.add_paragraph()
p.add_run('dhe, Daniel Henninger').bold = True
doc.add_paragraph('Wiesenstr. 28\n53773 Hennef')
doc.add_paragraph('(nachfolgend „Auftragnehmer" genannt)').italic = True

# Und
doc.add_heading('und', level=1)

# Zweite Partei (mit Platzhaltern)
p = doc.add_paragraph()
run = p.add_run('[KUNDENNAME]')
run.bold = True
# Gelbe Hervorhebung für Platzhalter
shading_elm = OxmlElement('w:shd')
shading_elm.set(qn('w:fill'), 'FFFF00')
run._r.get_or_add_rPr().append(shading_elm)

p = doc.add_paragraph()
run = p.add_run('[KUNDENADRESSE]\n[PLZ ORT]')
# Gelbe Hervorhebung
shading_elm = OxmlElement('w:shd')
shading_elm.set(qn('w:fill'), 'FFFF00')
run._r.get_or_add_rPr().append(shading_elm)

doc.add_paragraph('(nachfolgend „Auftraggeber" genannt)').italic = True

# Präambel
doc.add_heading('Präambel', level=1)
doc.add_paragraph(
    'Die Parteien beabsichtigen, im Bereich der Software "VoltMaster" zusammenzuarbeiten. '
    'Im Rahmen dieser Zusammenarbeit ist es erforderlich, dass die Parteien vertrauliche '
    'Informationen austauschen. Zum Schutz dieser vertraulichen Informationen vereinbaren '
    'die Parteien Folgendes:'
)

# § 1 Gegenstand der Vereinbarung
doc.add_heading('§ 1 Gegenstand der Vereinbarung', level=1)
doc.add_paragraph(
    '1. Gegenstand dieser Vereinbarung ist die Regelung des Umgangs mit vertraulichen '
    'Informationen, die im Rahmen der Zusammenarbeit bezüglich der Software "VoltMaster" '
    'zwischen den Parteien ausgetauscht werden.'
)
doc.add_paragraph('2. Die Zusammenarbeit umfasst insbesondere:')
doc.add_paragraph('• Entwicklung und Anpassung der Software VoltMaster', style='List Bullet')
doc.add_paragraph('• Beratung und Konzeption', style='List Bullet')
doc.add_paragraph('• Implementierung und Integration', style='List Bullet')
doc.add_paragraph('• Support und Wartung', style='List Bullet')

# § 2 Vertrauliche Informationen
doc.add_heading('§ 2 Vertrauliche Informationen', level=1)
doc.add_paragraph(
    '1. Als vertrauliche Informationen im Sinne dieser Vereinbarung gelten alle Informationen '
    'und Unterlagen, die eine Partei der anderen Partei zugänglich macht und die als vertraulich '
    'gekennzeichnet sind oder nach der Art der Information oder den Umständen der Übermittlung '
    'als vertraulich anzusehen sind.'
)
doc.add_paragraph('2. Insbesondere gelten als vertraulich:')
doc.add_paragraph('• Quellcode und Programmarchitektur der Software VoltMaster', style='List Bullet')
doc.add_paragraph('• Technische Dokumentationen und Spezifikationen', style='List Bullet')
doc.add_paragraph('• Geschäfts- und Betriebsgeheimnisse', style='List Bullet')
doc.add_paragraph('• Kundendaten und Kundenbeziehungen', style='List Bullet')
doc.add_paragraph('• Preisgestaltung und Geschäftsmodelle', style='List Bullet')
doc.add_paragraph('• Entwicklungspläne und Strategien', style='List Bullet')
doc.add_paragraph('• Sicherheitskonzepte und Zugangsdaten', style='List Bullet')

doc.add_paragraph('3. Nicht als vertraulich gelten Informationen, die:')
doc.add_paragraph('• zum Zeitpunkt der Übermittlung bereits öffentlich bekannt waren', style='List Bullet')
doc.add_paragraph('• nach der Übermittlung ohne Verschulden des Empfängers öffentlich bekannt werden', style='List Bullet')
doc.add_paragraph('• dem Empfänger bereits vor der Übermittlung bekannt waren', style='List Bullet')
doc.add_paragraph('• vom Empfänger unabhängig und ohne Verwendung vertraulicher Informationen entwickelt wurden', style='List Bullet')
doc.add_paragraph('• dem Empfänger rechtmäßig von Dritten ohne Geheimhaltungsverpflichtung zugänglich gemacht wurden', style='List Bullet')

# § 3 Geheimhaltungspflichten
doc.add_heading('§ 3 Geheimhaltungspflichten', level=1)
doc.add_paragraph(
    '1. Die Parteien verpflichten sich, vertrauliche Informationen streng geheim zu halten '
    'und nicht an Dritte weiterzugeben.'
)
doc.add_paragraph(
    '2. Die Parteien werden vertrauliche Informationen ausschließlich zum Zweck der '
    'vereinbarten Zusammenarbeit verwenden.'
)
doc.add_paragraph(
    '3. Die Parteien werden vertrauliche Informationen mit der gleichen Sorgfalt schützen, '
    'mit der sie eigene vertrauliche Informationen schützen, mindestens jedoch mit der '
    'Sorgfalt eines ordentlichen Kaufmanns.'
)
doc.add_paragraph(
    '4. Der Zugang zu vertraulichen Informationen wird auf diejenigen Mitarbeiter beschränkt, '
    'die diese für die Durchführung der Zusammenarbeit benötigen ("Need-to-know-Prinzip").'
)

# § 4 Erlaubte Offenlegung
doc.add_heading('§ 4 Erlaubte Offenlegung', level=1)
doc.add_paragraph('Eine Offenlegung vertraulicher Informationen ist zulässig, wenn:')
doc.add_paragraph('• der Offenlegende gesetzlich oder durch behördliche/gerichtliche Anordnung zur Offenlegung verpflichtet ist', style='List Bullet')
doc.add_paragraph('• die andere Partei der Offenlegung schriftlich zugestimmt hat', style='List Bullet')
doc.add_paragraph('Im Falle einer gesetzlichen Offenlegungspflicht ist die andere Partei unverzüglich zu informieren.')

# § 5 Rückgabe von Unterlagen
doc.add_heading('§ 5 Rückgabe von Unterlagen', level=1)
doc.add_paragraph(
    'Nach Beendigung der Zusammenarbeit oder auf Verlangen einer Partei sind alle überlassenen '
    'vertraulichen Informationen einschließlich aller Kopien zurückzugeben oder nachweislich zu vernichten.'
)

# § 6 Laufzeit
doc.add_heading('§ 6 Laufzeit', level=1)
doc.add_paragraph('1. Diese Vereinbarung tritt mit Unterzeichnung durch beide Parteien in Kraft.')
doc.add_paragraph('2. Die Geheimhaltungspflicht gilt für einen Zeitraum von ')
p = doc.paragraphs[-1]
p.add_run('5 Jahren').bold = True
p.add_run(' nach Beendigung der Zusammenarbeit.')

# § 7 Vertragsstrafe
doc.add_heading('§ 7 Vertragsstrafe', level=1)
p = doc.add_paragraph('Bei schuldhafter Verletzung der Geheimhaltungspflicht ist eine Vertragsstrafe in Höhe von ')
p.add_run('EUR 10.000,00').bold = True
p.add_run(' je Einzelfall fällig. Die Geltendmachung weitergehenden Schadensersatzes bleibt vorbehalten.')

# § 8 Schlussbestimmungen
doc.add_heading('§ 8 Schlussbestimmungen', level=1)
doc.add_paragraph('1. Änderungen und Ergänzungen dieser Vereinbarung bedürfen der Schriftform.')
doc.add_paragraph(
    '2. Sollten einzelne Bestimmungen dieser Vereinbarung unwirksam sein oder werden, '
    'bleibt die Wirksamkeit der übrigen Bestimmungen unberührt.'
)
doc.add_paragraph('3. Es gilt das Recht der Bundesrepublik Deutschland.')
doc.add_paragraph('4. Gerichtsstand ist Köln.')

# Unterschriften
doc.add_page_break()
doc.add_paragraph('\n\n')
doc.add_paragraph('Ort, Datum: _______________________')
doc.add_paragraph('\n\n\n')

# Tabelle für Unterschriften
table = doc.add_table(rows=1, cols=2)
left_cell = table.rows[0].cells[0]
right_cell = table.rows[0].cells[1]

left_cell.text = 'Für den Auftragnehmer:\n\n\n\n_______________________\ndhe, Daniel Henninger'
right_cell.text = 'Für den Auftraggeber:\n\n\n\n_______________________\n'
p = right_cell.add_paragraph()
run = p.add_run('[Name, Funktion]')
shading_elm = OxmlElement('w:shd')
shading_elm.set(qn('w:fill'), 'FFFF00')
run._r.get_or_add_rPr().append(shading_elm)

# Speichern
doc.save('docs/NDA_VoltMaster.docx')
print("NDA Word-Dokument wurde erstellt: docs/NDA_VoltMaster.docx")
